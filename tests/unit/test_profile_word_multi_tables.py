"""多表 Word：每个表对应 table_profile + builtin_prompt。"""

import os
import tempfile

from docx import Document

from src.core.profile import generate_profile_from_template


def test_word_multi_table_one_profile_and_builtin_prompt_per_table():
    doc = Document()
    doc.add_paragraph("甲表说明：经济指标")
    t1 = doc.add_table(rows=2, cols=2)
    t1.rows[0].cells[0].text = "城市"
    t1.rows[0].cells[1].text = "GDP"
    doc.add_paragraph("乙表说明：人口")
    t2 = doc.add_table(rows=2, cols=2)
    t2.rows[0].cells[0].text = "城市"
    t2.rows[0].cells[1].text = "人口"
    fd, path = tempfile.mkstemp(suffix=".docx")
    os.close(fd)
    try:
        doc.save(path)
        profile = generate_profile_from_template(template_path=path, mode="file")
        assert profile["template_mode"] == "word_multi_table"
        assert len(profile["table_profiles"]) == 2
        assert len(profile["table_specs"]) == 2
        for i, spec in enumerate(profile["table_specs"]):
            assert "table_profile" in spec
            assert "builtin_prompt" in spec
            assert spec["table_profile"]["table_index"] == i
            assert len(spec["table_profile"]["fields"]) >= 1
            assert "内置抽取提示" in spec["builtin_prompt"]
            assert profile["table_profiles"][i] == spec["table_profile"]
    finally:
        os.unlink(path)
