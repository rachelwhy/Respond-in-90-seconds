"""Word 多表模板检测：首行表头 + 表上方段落说明。"""

import os
import tempfile

from docx import Document

from src.core.template_detector import detect_template_structure


def test_single_table_word_template_unchanged():
    doc = Document()
    doc.add_paragraph("前言")
    t = doc.add_table(rows=2, cols=3)
    t.rows[0].cells[0].text = "列A"
    t.rows[0].cells[1].text = "列B"
    t.rows[0].cells[2].text = "列C"
    fd, path = tempfile.mkstemp(suffix=".docx")
    os.close(fd)
    try:
        doc.save(path)
        d = detect_template_structure(path)
        assert d["template_mode"] == "word_table"
        assert d["field_names"] == ["列A", "列B", "列C"]
    finally:
        os.unlink(path)


def test_multi_table_word_template_and_instruction_above():
    doc = Document()
    doc.add_paragraph("说明甲：填经济指标表")
    t1 = doc.add_table(rows=2, cols=2)
    t1.rows[0].cells[0].text = "城市"
    t1.rows[0].cells[1].text = "GDP（亿元）"
    doc.add_paragraph("说明乙：填人口表")
    t2 = doc.add_table(rows=2, cols=2)
    t2.rows[0].cells[0].text = "城市"
    t2.rows[0].cells[1].text = "人口（万人）"
    fd, path = tempfile.mkstemp(suffix=".docx")
    os.close(fd)
    try:
        doc.save(path)
        d = detect_template_structure(path)
        assert d["template_mode"] == "word_multi_table"
        assert len(d["table_specs"]) == 2
        assert d["table_specs"][0]["table_index"] == 0
        assert d["table_specs"][1]["table_index"] == 1
        assert "说明甲" in d["table_specs"][0]["instruction_above"]
        assert "说明乙" in d["table_specs"][1]["instruction_above"]
        assert d["table_specs"][0]["field_names"] == ["城市", "GDP（亿元）"]
        assert d["table_specs"][1]["field_names"] == ["城市", "人口（万人）"]
        # 并集字段：城市 只出现一次，其余按表顺序追加
        assert "城市" in d["field_names"]
        assert "GDP（亿元）" in d["field_names"]
        assert "人口（万人）" in d["field_names"]
    finally:
        os.unlink(path)
