"""按表切分输入 docx 片段。"""

import os
import tempfile

from docx import Document

from src.core.word_multi_segments import build_word_multi_table_segments, segments_from_input_docx


def test_segments_from_docx_two_tables():
    doc = Document()
    doc.add_paragraph("甲段说明")
    t1 = doc.add_table(rows=2, cols=2)
    t1.rows[0].cells[0].text = "A"
    t1.rows[0].cells[1].text = "B"
    doc.add_paragraph("乙段说明")
    t2 = doc.add_table(rows=2, cols=2)
    t2.rows[0].cells[0].text = "C"
    t2.rows[0].cells[1].text = "D"
    fd, path = tempfile.mkstemp(suffix=".docx")
    os.close(fd)
    try:
        doc.save(path)
        segs = segments_from_input_docx(path, 2)
        assert segs is not None and len(segs) == 2
        assert "甲段" in segs[0]
        assert "A" in segs[0] or "B" in segs[0]
        assert "乙段" in segs[1]
    finally:
        os.unlink(path)


def test_build_segments_matches_profile_table_count():
    profile = {
        "template_mode": "word_multi_table",
        "table_specs": [{"table_index": 0}, {"table_index": 1}],
    }
    doc = Document()
    doc.add_paragraph("p1")
    doc.add_table(rows=1, cols=1)
    doc.add_paragraph("p2")
    doc.add_table(rows=1, cols=1)
    fd, path = tempfile.mkstemp(suffix=".docx")
    os.close(fd)
    try:
        doc.save(path)
        documents = [{"path": path, "text": "x"}]
        segs = build_word_multi_table_segments(profile, "fallback全文", documents)
        assert len(segs) == 2
    finally:
        os.unlink(path)
