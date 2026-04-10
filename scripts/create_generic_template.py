#!/usr/bin/env python3
"""
创建通用Excel模板，用于无模板文件测试
"""

import os
from pathlib import Path
from openpyxl import Workbook

def create_generic_excel_template(output_path: Path = Path("data/template/generic_template.xlsx")):
    """创建通用Excel模板

    模板包含常见字段，适用于多种数据类型
    """
    # 确保目录存在
    output_path.parent.mkdir(parents=True, exist_ok=True)

    # 创建工作簿
    wb = Workbook()
    ws = wb.active
    ws.title = "数据模板"

    # 定义通用表头（包含常见字段）
    headers = [
        "序号", "名称", "类型", "数值", "单位", "日期",
        "地点", "城市", "省份", "国家", "描述", "备注"
    ]

    # 写入表头
    for col_idx, header in enumerate(headers, 1):
        ws.cell(row=1, column=col_idx, value=header)

    # 添加示例数据行（用于说明）
    example_data = [
        [1, "示例项目", "经济", 1000, "万元", "2025-01-01",
         "示例地点", "北京", "北京", "中国", "示例描述", "示例备注"],
        [2, "测试数据", "环境", 50.5, "μg/m³", "2025-02-01",
         "测试站点", "上海", "上海", "中国", "测试描述", "测试备注"]
    ]

    for row_idx, row_data in enumerate(example_data, 2):
        for col_idx, value in enumerate(row_data, 1):
            ws.cell(row=row_idx, column=col_idx, value=value)

    # 调整列宽
    for col in ws.columns:
        max_length = 0
        column = col[0].column_letter
        for cell in col:
            try:
                if len(str(cell.value)) > max_length:
                    max_length = len(str(cell.value))
            except:
                pass
        adjusted_width = min(max_length + 2, 50)
        ws.column_dimensions[column].width = adjusted_width

    # 保存文件
    wb.save(output_path)
    print(f"通用Excel模板已创建: {output_path}")
    print(f"包含字段: {', '.join(headers)}")

    return output_path


def create_generic_word_template(output_path: Path = Path("data/template/generic_template.docx")):
    """创建通用Word模板

    包含一个简单表格，适用于Word文档测试
    """
    try:
        from docx import Document
        from docx.shared import Inches

        # 确保目录存在
        output_path.parent.mkdir(parents=True, exist_ok=True)

        # 创建文档
        doc = Document()

        # 添加标题
        doc.add_heading('通用数据模板', 0)
        doc.add_paragraph('这是一个通用Word模板，用于数据提取测试。')

        # 添加表格
        table = doc.add_table(rows=3, cols=5)
        table.style = 'Table Grid'

        # 设置表头
        headers = ['项目', '数值', '单位', '日期', '备注']
        for i, header in enumerate(headers):
            table.cell(0, i).text = header

        # 设置示例数据
        examples = [
            ['GDP', '10000', '亿元', '2025年', '示例数据'],
            ['人口', '1500', '万人', '2025年', '示例数据'],
            ['PM2.5', '35', 'μg/m³', '2025年', '示例数据']
        ]

        for row_idx, example in enumerate(examples, 1):
            for col_idx, value in enumerate(example):
                table.cell(row_idx, col_idx).text = str(value)

        # 保存文档
        doc.save(output_path)
        print(f"通用Word模板已创建: {output_path}")

        return output_path

    except ImportError:
        print("⚠ python-docx未安装，跳过Word模板创建")
        return None


if __name__ == '__main__':
    print("创建通用模板文件...")

    # 创建Excel模板
    excel_template = create_generic_excel_template()

    # 创建Word模板
    word_template = create_generic_word_template()

    print("\n模板创建完成。")
    print(f"Excel模板: {excel_template}")
    print(f"Word模板: {word_template or '未创建'}")

    # 创建模板映射文件
    template_map = {
        "excel": str(excel_template),
        "word": str(word_template) if word_template else str(excel_template),
        "markdown": str(excel_template),
        "text": str(excel_template),
        "mixed": str(excel_template)
    }

    # 保存模板映射
    map_path = Path("data/template/template_mapping.json")
    import json
    with open(map_path, 'w', encoding='utf-8') as f:
        json.dump(template_map, f, ensure_ascii=False, indent=2)

    print(f"模板映射已保存: {map_path}")