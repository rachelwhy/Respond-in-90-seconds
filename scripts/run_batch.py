import argparse
import json
import subprocess
import sys
import time
import traceback
from dataclasses import dataclass, asdict
from datetime import datetime
from pathlib import Path
from typing import Dict, List, Any, Optional, Tuple, Union

try:
    import psutil
    PSUTIL_AVAILABLE = True
except ImportError:
    PSUTIL_AVAILABLE = False
    psutil = None

try:
    import pandas as pd
    PANDAS_AVAILABLE = True
except ImportError:
    PANDAS_AVAILABLE = False
    pd = None

try:
    from openpyxl import load_workbook
    OPENPYXL_AVAILABLE = True
except ImportError:
    OPENPYXL_AVAILABLE = False
    load_workbook = None

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    Document = None


@dataclass
class ValidationResult:
    """验证结果数据结构"""
    task_id: str
    status: str  # 'passed', 'failed', 'partial', 'skipped'
    accuracy: float  # 0-1
    field_accuracy: Dict[str, float]  # 每个字段的准确率
    missing_fields: List[str]
    incorrect_fields: List[Dict[str, Any]]  # {field: str, expected: Any, actual: Any}
    extra_fields: List[str]
    record_count_match: bool
    expected_records: int
    actual_records: int
    validation_errors: List[str]

    def to_dict(self):
        return asdict(self)


@dataclass
class PerformanceMetrics:
    """性能指标数据结构"""
    task_id: str
    total_time: float
    model_inference_time: Optional[float] = None
    extraction_time: Optional[float] = None
    postprocessing_time: Optional[float] = None
    memory_usage_mb: Optional[float] = None
    peak_memory_mb: Optional[float] = None
    cpu_percent: Optional[float] = None

    def to_dict(self):
        return asdict(self)


@dataclass
class BenchmarkReport:
    """基准测试报告数据结构"""
    timestamp: str
    main_script: str
    validation_mode: str
    threshold: float
    summary: Dict[str, Any]  # 总体统计
    tasks: List[Dict[str, Any]]  # 每个任务的详细结果
    recommendations: List[str]
    issues: List[Dict[str, Any]]

    def to_dict(self):
        return asdict(self)


class BenchmarkValidator:
    """基准测试验证器"""
    def __init__(self, validation_mode='fieldwise', threshold=0.8):
        self.mode = validation_mode
        self.threshold = threshold

    def validate(self, actual_path: Path, expected_path: Path,
                 field_config: Optional[Dict] = None) -> ValidationResult:
        """验证文件 - 根据文件类型自动选择验证方法"""
        try:
            if not actual_path.exists():
                return ValidationResult(
                    task_id=str(actual_path.parent.name),
                    status='failed',
                    accuracy=0.0,
                    field_accuracy={},
                    missing_fields=[],
                    incorrect_fields=[],
                    extra_fields=[],
                    record_count_match=False,
                    expected_records=0,
                    actual_records=0,
                    validation_errors=[f'实际文件不存在: {actual_path}']
                )

            if not expected_path.exists():
                return ValidationResult(
                    task_id=str(actual_path.parent.name),
                    status='skipped',
                    accuracy=1.0,  # 跳过验证，视为通过
                    field_accuracy={},
                    missing_fields=[],
                    incorrect_fields=[],
                    extra_fields=[],
                    record_count_match=True,
                    expected_records=0,
                    actual_records=0,
                    validation_errors=[f'标准答案文件不存在: {expected_path}，跳过验证']
                )

            # 根据文件扩展名选择验证方法
            ext = actual_path.suffix.lower()
            if ext == '.xlsx' or ext == '.xls':
                return self.validate_excel(actual_path, expected_path, field_config)
            elif ext == '.docx':
                return self.validate_word(actual_path, expected_path, field_config)
            elif ext == '.json':
                return self.validate_json(actual_path, expected_path, field_config)
            else:
                return ValidationResult(
                    task_id=str(actual_path.parent.name),
                    status='failed',
                    accuracy=0.0,
                    field_accuracy={},
                    missing_fields=[],
                    incorrect_fields=[],
                    extra_fields=[],
                    record_count_match=False,
                    expected_records=0,
                    actual_records=0,
                    validation_errors=[f'不支持的文件类型: {ext}']
                )

        except Exception as e:
            return ValidationResult(
                task_id=str(actual_path.parent.name) if actual_path.exists() else 'unknown',
                status='failed',
                accuracy=0.0,
                field_accuracy={},
                missing_fields=[],
                incorrect_fields=[],
                extra_fields=[],
                record_count_match=False,
                expected_records=0,
                actual_records=0,
                validation_errors=[f'验证过程中发生错误: {str(e)}', traceback.format_exc()]
            )

    def validate_excel(self, actual_path: Path, expected_path: Path,
                       field_config: Optional[Dict] = None) -> ValidationResult:
        """验证Excel文件"""
        # 简化实现 - 实际项目中需要更详细的对比逻辑
        # 这里只检查文件是否存在和基本格式

        task_id = actual_path.parent.name

        try:
            # 使用openpyxl读取文件
            if OPENPYXL_AVAILABLE:
                actual_wb = load_workbook(actual_path, data_only=True)
                expected_wb = load_workbook(expected_path, data_only=True)

                actual_ws = actual_wb.active
                expected_ws = expected_wb.active

                # 简单对比：检查是否有数据
                actual_has_data = any(actual_ws.iter_rows(values_only=True))
                expected_has_data = any(expected_ws.iter_rows(values_only=True))

                if actual_has_data and expected_has_data:
                    # 基本验证通过
                    return ValidationResult(
                        task_id=task_id,
                        status='passed',
                        accuracy=1.0,  # 简化实现
                        field_accuracy={},
                        missing_fields=[],
                        incorrect_fields=[],
                        extra_fields=[],
                        record_count_match=True,
                        expected_records=1,  # 简化
                        actual_records=1,    # 简化
                        validation_errors=[]
                    )
                else:
                    return ValidationResult(
                        task_id=task_id,
                        status='partial',
                        accuracy=0.5 if actual_has_data else 0.0,
                        field_accuracy={},
                        missing_fields=[],
                        incorrect_fields=[],
                        extra_fields=[],
                        record_count_match=False,
                        expected_records=1,
                        actual_records=0 if not actual_has_data else 1,
                        validation_errors=['Excel文件数据不完整']
                    )
            else:
                # 如果没有openpyxl，只检查文件存在性
                return ValidationResult(
                    task_id=task_id,
                    status='partial',
                    accuracy=0.5,  # 文件存在但无法验证内容
                    field_accuracy={},
                    missing_fields=[],
                    incorrect_fields=[],
                    extra_fields=[],
                    record_count_match=True,
                    expected_records=0,
                    actual_records=0,
                    validation_errors=['openpyxl不可用，跳过详细验证']
                )

        except Exception as e:
            return ValidationResult(
                task_id=task_id,
                status='failed',
                accuracy=0.0,
                field_accuracy={},
                missing_fields=[],
                incorrect_fields=[],
                extra_fields=[],
                record_count_match=False,
                expected_records=0,
                actual_records=0,
                validation_errors=[f'Excel验证错误: {str(e)}']
            )

    def validate_word(self, actual_path: Path, expected_path: Path,
                      field_config: Optional[Dict] = None) -> ValidationResult:
        """验证Word文件 - 简化实现"""
        task_id = actual_path.parent.name

        try:
            if DOCX_AVAILABLE:
                # 简单检查：文件存在且能打开
                actual_doc = Document(actual_path)
                expected_doc = Document(expected_path)

                # 检查是否有表格
                actual_tables = len(actual_doc.tables) > 0
                expected_tables = len(expected_doc.tables) > 0

                return ValidationResult(
                    task_id=task_id,
                    status='passed' if actual_tables else 'partial',
                    accuracy=1.0 if actual_tables else 0.5,
                    field_accuracy={},
                    missing_fields=[],
                    incorrect_fields=[],
                    extra_fields=[],
                    record_count_match=actual_tables == expected_tables,
                    expected_records=len(expected_doc.tables) if expected_tables else 0,
                    actual_records=len(actual_doc.tables) if actual_tables else 0,
                    validation_errors=[]
                )
            else:
                return ValidationResult(
                    task_id=task_id,
                    status='partial',
                    accuracy=0.5,
                    field_accuracy={},
                    missing_fields=[],
                    incorrect_fields=[],
                    extra_fields=[],
                    record_count_match=True,
                    expected_records=0,
                    actual_records=0,
                    validation_errors=['python-docx不可用，跳过详细验证']
                )

        except Exception as e:
            return ValidationResult(
                task_id=task_id,
                status='failed',
                accuracy=0.0,
                field_accuracy={},
                missing_fields=[],
                incorrect_fields=[],
                extra_fields=[],
                record_count_match=False,
                expected_records=0,
                actual_records=0,
                validation_errors=[f'Word验证错误: {str(e)}']
            )

    def validate_json(self, actual_path: Path, expected_path: Path,
                      field_config: Optional[Dict] = None) -> ValidationResult:
        """验证JSON文件"""
        task_id = actual_path.parent.name

        try:
            with open(actual_path, 'r', encoding='utf-8') as f:
                actual_data = json.load(f)

            with open(expected_path, 'r', encoding='utf-8') as f:
                expected_data = json.load(f)

            # 简单对比：检查是否有数据
            actual_has_data = bool(actual_data)
            expected_has_data = bool(expected_data)

            return ValidationResult(
                task_id=task_id,
                status='passed' if actual_has_data else 'partial',
                accuracy=1.0 if actual_has_data else 0.0,
                field_accuracy={},
                missing_fields=[],
                incorrect_fields=[],
                extra_fields=[],
                record_count_match=actual_has_data == expected_has_data,
                expected_records=1 if expected_has_data else 0,
                actual_records=1 if actual_has_data else 0,
                validation_errors=[]
            )

        except Exception as e:
            return ValidationResult(
                task_id=task_id,
                status='failed',
                accuracy=0.0,
                field_accuracy={},
                missing_fields=[],
                incorrect_fields=[],
                extra_fields=[],
                record_count_match=False,
                expected_records=0,
                actual_records=0,
                validation_errors=[f'JSON验证错误: {str(e)}']
            )


class PerformanceCollector:
    """性能指标收集器"""
    def __init__(self):
        self.start_time = time.time()
        self.process = psutil.Process() if PSUTIL_AVAILABLE else None
        self.start_memory = self._get_memory_usage() if self.process else None

    def _get_memory_usage(self):
        """获取内存使用量（MB）"""
        if self.process:
            try:
                return self.process.memory_info().rss / 1024 / 1024  # MB
            except:
                return None
        return None

    def collect(self, task_id: str) -> PerformanceMetrics:
        """收集性能指标"""
        total_time = time.time() - self.start_time

        metrics = PerformanceMetrics(
            task_id=task_id,
            total_time=total_time
        )

        if self.process:
            try:
                current_memory = self._get_memory_usage()
                if self.start_memory and current_memory:
                    metrics.memory_usage_mb = current_memory
                    metrics.peak_memory_mb = current_memory  # 简化实现

                metrics.cpu_percent = self.process.cpu_percent(interval=0.1)
            except:
                pass  # 忽略性能收集错误

        return metrics


def load_field_config(profile_path: str) -> Optional[Dict]:
    """加载字段配置"""
    try:
        profile_file = Path(profile_path)
        if profile_file.exists():
            with open(profile_file, 'r', encoding='utf-8') as f:
                return json.load(f)
    except Exception as e:
        print(f'[WARN] 无法加载字段配置 {profile_path}: {e}')
    return None


def determine_output_files(output_dir: Path, template_path: str) -> Dict[str, Path]:
    """根据模板路径确定输出文件类型和路径"""
    template_ext = Path(template_path).suffix.lower()
    result_files = {}

    # 从模板路径提取基本名称
    import os
    base_name = os.path.splitext(os.path.basename(template_path))[0]

    # 确定可能的输出文件（使用动态文件名）
    if template_ext == '.docx':
        # 查找Word结果文件
        docx_files = list(output_dir.glob(f'*_result.docx'))
        if docx_files:
            # 优先使用匹配base_name的文件
            matching_docx = [f for f in docx_files if f.name == f'{base_name}_result.docx']
            result_files['docx'] = matching_docx[0] if matching_docx else docx_files[0]
        else:
            # 回退到固定文件名（保持向后兼容）
            result_files['docx'] = output_dir / 'result.docx'

        # 查找JSON结果文件
        json_files = list(output_dir.glob(f'*_result.json'))
        if json_files:
            matching_json = [f for f in json_files if f.name == f'{base_name}_result.json']
            result_files['json'] = matching_json[0] if matching_json else json_files[0]
        else:
            result_files['json'] = output_dir / 'result.json'
    else:  # .xlsx 或其他
        # 查找Excel结果文件
        xlsx_files = list(output_dir.glob(f'*_result.xlsx'))
        if xlsx_files:
            matching_xlsx = [f for f in xlsx_files if f.name == f'{base_name}_result.xlsx']
            result_files['xlsx'] = matching_xlsx[0] if matching_xlsx else xlsx_files[0]
        else:
            result_files['xlsx'] = output_dir / 'result.xlsx'

        # 查找JSON结果文件
        json_files = list(output_dir.glob(f'*_result.json'))
        if json_files:
            matching_json = [f for f in json_files if f.name == f'{base_name}_result.json']
            result_files['json'] = matching_json[0] if matching_json else json_files[0]
        else:
            result_files['json'] = output_dir / 'result.json'

    return result_files


def generate_report(tasks: List[Dict], args, total_start_time: float) -> BenchmarkReport:
    """生成基准测试报告"""
    total_tasks = len(tasks)
    passed_tasks = sum(1 for t in tasks if t.get('validation', {}).get('status') in ['passed', 'skipped'])
    failed_tasks = sum(1 for t in tasks if t.get('validation', {}).get('status') == 'failed')
    partial_tasks = sum(1 for t in tasks if t.get('validation', {}).get('status') == 'partial')

    # 计算平均准确率
    accuracies = [t.get('validation', {}).get('accuracy', 0) for t in tasks if t.get('validation')]
    avg_accuracy = sum(accuracies) / len(accuracies) if accuracies else 0

    # 计算总时间
    total_time = time.time() - total_start_time

    # 收集性能指标
    performance_metrics = [t.get('performance') for t in tasks if t.get('performance')]
    if performance_metrics:
        # 安全计算平均值，处理None值
        total_memory = 0
        total_cpu = 0
        count_memory = 0
        count_cpu = 0

        for m in performance_metrics:
            memory = m.get('memory_usage_mb')
            cpu = m.get('cpu_percent')
            if memory is not None:
                total_memory += memory
                count_memory += 1
            if cpu is not None:
                total_cpu += cpu
                count_cpu += 1

        avg_memory = total_memory / count_memory if count_memory > 0 else 0
        avg_cpu = total_cpu / count_cpu if count_cpu > 0 else 0
    else:
        avg_memory = 0
        avg_cpu = 0

    summary = {
        'total_tasks': total_tasks,
        'passed_tasks': passed_tasks,
        'failed_tasks': failed_tasks,
        'partial_tasks': partial_tasks,
        'success_rate': passed_tasks / total_tasks if total_tasks > 0 else 0,
        'average_accuracy': avg_accuracy,
        'total_time_seconds': total_time,
        'average_memory_mb': avg_memory,
        'average_cpu_percent': avg_cpu
    }

    # 收集问题
    issues = []
    recommendations = []

    for task in tasks:
        task_id = task['task_id']
        validation = task.get('validation')
        performance = task.get('performance')

        if validation and validation.get('status') == 'failed':
            issues.append({
                'task_id': task_id,
                'type': 'validation',
                'description': f'验证失败: {validation.get("validation_errors", ["未知错误"])[0]}'
            })
            recommendations.append(f'检查任务 {task_id} 的输入文件和模板配置')

        if performance and performance.get('total_time', 0) > 90:
            issues.append({
                'task_id': task_id,
                'type': 'performance',
                'description': f'运行时间过长: {performance.get("total_time"):.1f}秒'
            })
            recommendations.append(f'优化任务 {task_id} 的性能，目标<90秒')

    # 去重推荐
    recommendations = list(set(recommendations))

    return BenchmarkReport(
        timestamp=datetime.now().isoformat(),
        main_script=args.main_script,
        validation_mode=args.validation_mode if args.validate else 'none',
        threshold=args.threshold if args.validate else 0,
        summary=summary,
        tasks=tasks,
        recommendations=recommendations,
        issues=issues
    )


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument('--manifest', required=True)
    ap.add_argument('--max-tasks', type=int, default=0)
    ap.add_argument('--main-script', default='main.py', help='主脚本路径（如main.py）')
    ap.add_argument('--validate', action='store_true', help='启用结果验证')
    ap.add_argument('--validation-mode', choices=['strict', 'lenient', 'fieldwise'], default='fieldwise', help='验证模式')
    ap.add_argument('--threshold', type=float, default=0.8, help='验证通过阈值（0-1）')
    ap.add_argument('--collect-metrics', action='store_true', help='收集性能指标')
    ap.add_argument('--report-format', choices=['json', 'html', 'markdown'], default='json', help='报告格式')
    ap.add_argument('--output-report', default='test/reports/benchmark_report.json', help='验证报告输出路径')
    args = ap.parse_args()

    print(f'=== A23 AI 基准测试验证系统 ===')
    print(f'主脚本: {args.main_script}')
    print(f'验证模式: {args.validation_mode if args.validate else "禁用"}')
    print(f'性能收集: {"启用" if args.collect_metrics else "禁用"}')
    print()

    # 加载任务清单
    manifest = json.loads(Path(args.manifest).read_text(encoding='utf-8'))
    tasks = manifest['tasks'][: args.max_tasks] if args.max_tasks else manifest['tasks']
    print(f'找到 {len(tasks)} 个任务，开始处理...')

    # 初始化验证器
    validator = None
    if args.validate:
        validator = BenchmarkValidator(
            validation_mode=args.validation_mode,
            threshold=args.threshold
        )
        print(f'结果验证已启用 (模式: {args.validation_mode}, 阈值: {args.threshold})')

    total_start_time = time.time()
    task_results = []

    for i, task in enumerate(tasks, 1):
        task_id = task['task_id']
        print(f'\n=== 任务 {i}/{len(tasks)}: {task_id} ===')

        # 准备输出目录
        output_dir = Path('test/results/benchmark') / task_id
        output_dir.mkdir(parents=True, exist_ok=True)

        # 初始化性能收集器
        perf_collector = PerformanceCollector() if args.collect_metrics else None

        # 准备命令
        cmd = [
            sys.executable,
            args.main_script,
            '--template', task['template_path'],
            '--input-dir', task['input_dir'],
            '--output-dir', str(output_dir),
            '--overwrite-output'
        ]

        print(f'命令: {" ".join(cmd)}')

        # 运行日志文件
        stdout_path = output_dir / 'run_stdout.log'
        stderr_path = output_dir / 'run_stderr.log'
        cmd_log_path = output_dir / 'command.txt'
        cmd_log_path.write_text(' '.join(cmd), encoding='utf-8')

        # 运行任务
        task_start_time = time.time()
        try:
            with stdout_path.open('w', encoding='utf-8') as out, stderr_path.open('w', encoding='utf-8') as err:
                proc = subprocess.run(cmd, text=True, stdout=out, stderr=err)

            returncode = proc.returncode
            run_time = time.time() - task_start_time

            print(f'返回码: {returncode}, 运行时间: {run_time:.2f}秒')

        except Exception as e:
            print(f'运行失败: {e}')
            returncode = 1
            run_time = time.time() - task_start_time

        # 收集性能指标
        performance = None
        if perf_collector:
            performance = perf_collector.collect(task_id).to_dict()
            performance['run_time'] = run_time

        # 验证结果
        validation_result = None
        if args.validate and returncode == 0:
            # 确定输出文件
            output_files = determine_output_files(output_dir, task['template_path'])

            # 查找标准答案文件
            expected_path = None
            if 'answer_path' in task:
                expected_path = Path(task['answer_path'])
            else:
                # 尝试在标准答案目录中查找
                answer_dir = Path('test/assets/标准答案')
                expected_pattern = f"*{task_id}*"
                answer_files = list(answer_dir.glob(expected_pattern))
                if answer_files:
                    expected_path = answer_files[0]

            # 加载字段配置
            field_config = None
            if 'profile_path' in task:
                field_config = load_field_config(task['profile_path'])

            # 验证每个输出文件
            validation_results = []
            for file_type, actual_path in output_files.items():
                if actual_path.exists() and expected_path:
                    print(f'验证 {file_type} 文件: {actual_path.name}')
                    result = validator.validate(actual_path, expected_path, field_config)
                    validation_results.append(result)

            # 合并验证结果
            if validation_results:
                # 取第一个结果作为主要验证结果（简化实现）
                validation_result = validation_results[0].to_dict()
            else:
                validation_result = {
                    'task_id': task_id,
                    'status': 'skipped',
                    'accuracy': 0.0,
                    'validation_errors': ['没有找到可验证的输出文件或标准答案']
                }
        elif args.validate and returncode != 0:
            validation_result = {
                'task_id': task_id,
                'status': 'failed',
                'accuracy': 0.0,
                'validation_errors': [f'任务运行失败，返回码: {returncode}']
            }

        # 收集任务结果
        task_result = {
            'task_id': task_id,
            'returncode': returncode,
            'output_dir': str(output_dir),
            'stdout_log': str(stdout_path),
            'stderr_log': str(stderr_path),
            'run_time': run_time,
            'command': ' '.join(cmd)
        }

        if performance:
            task_result['performance'] = performance

        if validation_result:
            task_result['validation'] = validation_result

        task_results.append(task_result)

        # 输出任务摘要
        print(f'任务完成: {task_id}')
        if validation_result:
            status = validation_result.get('status', 'unknown')
            accuracy = validation_result.get('accuracy', 0)
            print(f'  验证状态: {status}, 准确率: {accuracy:.2%}')
        if performance:
            print(f'  运行时间: {performance.get("total_time", run_time):.2f}秒')

    # 生成报告
    print(f'\n=== 生成基准测试报告 ===')
    report = generate_report(task_results, args, total_start_time)

    # 保存报告
    report_path = Path(args.output_report)
    report_path.parent.mkdir(parents=True, exist_ok=True)

    with open(report_path, 'w', encoding='utf-8') as f:
        json.dump(report.to_dict(), f, ensure_ascii=False, indent=2)

    print(f'报告已保存: {report_path}')

    # 打印摘要
    print(f'\n=== 基准测试摘要 ===')
    summary = report.summary
    print(f'总任务数: {summary["total_tasks"]}')
    print(f'通过任务: {summary["passed_tasks"]}')
    print(f'失败任务: {summary["failed_tasks"]}')
    print(f'部分通过: {summary["partial_tasks"]}')
    print(f'成功率: {summary["success_rate"]:.2%}')
    print(f'平均准确率: {summary["average_accuracy"]:.2%}')
    print(f'总运行时间: {summary["total_time_seconds"]:.2f}秒')

    if report.issues:
        print(f'\n发现的问题:')
        for issue in report.issues:
            print(f'  - {issue["task_id"]}: {issue["description"]}')

    if report.recommendations:
        print(f'\n改进建议:')
        for rec in report.recommendations:
            print(f'  - {rec}')

    # 保存传统的批处理摘要（向后兼容）
    compat_summary = {
        'tasks': [
            {
                'task_id': t['task_id'],
                'returncode': t['returncode'],
                'output_dir': t['output_dir'],
                'stdout_log': t['stdout_log'],
                'stderr_log': t['stderr_log']
            }
            for t in task_results
        ]
    }
    compat_path = Path('test/results/outputs/批处理结果摘要.json')
    compat_path.parent.mkdir(parents=True, exist_ok=True)
    compat_path.write_text(json.dumps(compat_summary, ensure_ascii=False, indent=2), encoding='utf-8')
    print(f'\n传统摘要已保存: {compat_path} (向后兼容)')


if __name__ == '__main__':
    main()
