"""
文档加载模块：支持多种文件格式解析
支持格式：Excel、Word、Markdown、TXT、PDF
返回统一的结构化数据格式
"""

import os
import re
import pandas as pd
from docx import Document
from typing import Dict, Any, Optional, Tuple, List


class DocumentLoader:
    """
    文档加载器：支持Excel、Word、Markdown、TXT、PDF
    返回统一的结构化数据格式
    """

    def __init__(self):
        self.supported_formats = ['.docx', '.doc', '.xlsx', '.xls', '.md', '.txt', '.pdf']
        print(f"📄 文档加载器初始化，支持格式: {', '.join(self.supported_formats)}")

    def load(self, path: str, filename: str) -> Dict[str, Any]:
        """
        统一加载接口
        参数：
            path: 文件路径
            filename: 文件名
        返回：
            结构化文档数据
        """
        ext = os.path.splitext(filename)[1].lower()
        result = {
            "filename": filename,
            "path": path,
            "ext": ext,
            "load_time": None
        }

        import time
        start_time = time.time()

        try:
            if ext in ['.docx', '.doc']:
                data, err = self._load_docx(path)
            elif ext in ['.xlsx', '.xls']:
                data, err = self._load_excel(path)
            elif ext in ['.md']:
                data, err = self._load_md(path)
            elif ext in ['.txt']:
                data, err = self._load_txt(path)
            elif ext in ['.pdf']:
                data, err = self._load_pdf(path)
            else:
                result["error"] = f"不支持的文件类型: {ext}"
                return result

            if err:
                result["error"] = err
            else:
                result.update(data)

        except Exception as e:
            result["error"] = f"加载失败: {str(e)}"

        result["load_time"] = round(time.time() - start_time, 3)
        return result

    # ========== Excel 加载（结构化输出） ==========
    def _load_excel(self, path: str) -> Tuple[Dict, Optional[str]]:
        """
        加载Excel文件，返回结构化表格数据
        输出格式符合魏嘉华要求：
        {
            "type": "excel",
            "file_name": "xxx.xlsx",
            "sheets": [
                {
                    "sheet_name": "Sheet1",
                    "headers": ["列名1", "列名2", ...],
                    "rows": [
                        {
                            "row_index": 2,
                            "values": {
                                "列名1": "值1",
                                "列名2": "值2",
                                ...
                            },
                            "raw_cells": {
                                "A2": "值1",
                                "B2": "值2",
                                ...
                            }
                        }
                    ]
                }
            ]
        }
        """
        try:
            # 读取所有sheet
            excel_file = pd.read_excel(path, sheet_name=None, dtype=str)  # 全部转成字符串，避免NaN问题

            sheets = []

            for sheet_name, df in excel_file.items():
                # 获取表头
                headers = df.columns.tolist()

                # 处理数据行
                rows = []
                for idx, row in df.iterrows():
                    values = {}
                    raw_cells = {}

                    for col_idx, col_name in enumerate(headers):
                        value = row[col_name]
                        # 处理NaN和None
                        if pd.isna(value) or value is None:
                            value = ""
                        else:
                            value = str(value).strip()

                        values[col_name] = value

                        # 构建单元格坐标（A、B、C...Z, AA, AB...）
                        col_letter = self._col_idx_to_letter(col_idx)
                        raw_cells[f"{col_letter}{idx + 2}"] = value

                    rows.append({
                        "row_index": idx + 2,  # Excel行号从2开始（第1行是表头）
                        "values": values,
                        "raw_cells": raw_cells
                    })

                sheets.append({
                    "sheet_name": sheet_name,
                    "headers": headers,
                    "rows": rows
                })

            result = {
                "type": "excel",
                "file_name": os.path.basename(path),
                "sheets": sheets,
                "sheet_count": len(sheets),
                "total_rows": sum(len(s["rows"]) for s in sheets)
            }

            return result, None

        except Exception as e:
            return None, f"Excel解析失败: {str(e)}"

    def _col_idx_to_letter(self, idx: int) -> str:
        """将列索引转换为Excel列字母（0->A, 1->B, 25->Z, 26->AA, 27->AB...）"""
        letter = ""
        idx += 1  # 转为1-based
        while idx > 0:
            idx -= 1
            letter = chr(idx % 26 + 65) + letter
            idx //= 26
        return letter

    # ========== Word 加载 ==========
    def _load_docx(self, path: str) -> Tuple[Dict, Optional[str]]:
        """
        加载Word文档，返回结构化信息
        输出格式：
        {
            "type": "word",
            "paragraphs": [...],
            "tables": [markdown_table, ...],
            "lists": [...],
            "titles": [(level, title), ...],
            "text": "完整文本"
        }
        """
        result = {
            "type": "word",
            "paragraphs": [],
            "tables": [],
            "lists": [],
            "titles": [],
            "text": ""
        }

        try:
            doc = Document(path)

            # 段落
            for para in doc.paragraphs:
                if para.text.strip():
                    result["paragraphs"].append(para.text)
                    # 识别标题
                    if para.style.name.startswith('Heading'):
                        level = int(para.style.name.replace('Heading', '')) if para.style.name != 'Heading' else 1
                        result["titles"].append((level, para.text))

            # 表格 - 转为markdown格式
            for table in doc.tables:
                table_md = self._table_to_markdown_from_docx(table)
                if table_md:
                    result["tables"].append(table_md)

            # 列表 - 识别以项目符号开头的段落
            for para in doc.paragraphs:
                if para.text.strip() and (para.style.name.startswith('List') or
                                          para.text.startswith(('•', '-', '*', '·'))):
                    result["lists"].append(para.text)

            result["text"] = "\n".join(result["paragraphs"])
            return result, None

        except Exception as e:
            return None, f"Word解析失败: {str(e)}"

    def _table_to_markdown_from_docx(self, table) -> str:
        """将python-docx表格转换为Markdown格式"""
        try:
            rows = []
            for i, row in enumerate(table.rows):
                cells = []
                for cell in row.cells:
                    # 获取单元格文本
                    cell_text = cell.text.strip().replace('\n', ' ')
                    cells.append(cell_text)

                if i == 0:
                    # 表头
                    rows.append("| " + " | ".join(cells) + " |")
                    rows.append("|" + "|".join([" --- "] * len(cells)) + "|")
                else:
                    # 数据行
                    rows.append("| " + " | ".join(cells) + " |")

            return "\n".join(rows)
        except:
            return ""

    # ========== Markdown 加载 ==========
    def _load_md(self, path: str) -> Tuple[Dict, Optional[str]]:
        """
        加载Markdown文档，返回结构化信息
        输出格式：
        {
            "type": "markdown",
            "titles": [(level, title), ...],
            "tables": [markdown_table, ...],
            "lists": [list_item, ...],
            "paragraphs": [paragraph, ...],
            "text": "完整文本"
        }
        """
        result = {
            "type": "markdown",
            "titles": [],
            "tables": [],
            "lists": [],
            "paragraphs": [],
            "text": ""
        }

        try:
            with open(path, 'r', encoding='utf-8') as f:
                content = f.read()

            lines = content.split('\n')
            in_code_block = False
            current_para = []
            in_table = False
            table_lines = []

            for i, line in enumerate(lines):
                # 跳过代码块
                if line.startswith('```'):
                    in_code_block = not in_code_block
                    continue
                if in_code_block:
                    continue

                # 标题
                title_match = re.match(r'^(#{1,6})\s+(.+)$', line)
                if title_match:
                    level = len(title_match.group(1))
                    title = title_match.group(2)
                    result["titles"].append((level, title))
                    continue

                # 表格行
                if re.match(r'^\|.+\|$', line):
                    if not in_table:
                        in_table = True
                        table_lines = [line]
                    else:
                        table_lines.append(line)
                    continue
                else:
                    if in_table and table_lines:
                        # 表格结束，保存表格
                        result["tables"].append("\n".join(table_lines))
                        in_table = False
                        table_lines = []

                # 列表项
                list_match = re.match(r'^[\*\-\+]\s+(.+)$', line) or \
                             re.match(r'^\d+\.\s+(.+)$', line)
                if list_match:
                    result["lists"].append(list_match.group(1))
                    continue

                # 普通段落
                if line.strip():
                    current_para.append(line.strip())
                else:
                    if current_para:
                        result["paragraphs"].append(" ".join(current_para))
                        current_para = []

            # 处理最后一个表格
            if in_table and table_lines:
                result["tables"].append("\n".join(table_lines))

            # 处理最后一个段落
            if current_para:
                result["paragraphs"].append(" ".join(current_para))

            # 提取表格（正则表达式兜底）
            tables = re.findall(r'(\|.+\|\r?\n\|[-|\s]+\|\r?\n(?:\|.+\|\r?\n)+)', content)
            for table in tables:
                if table.strip() not in result["tables"]:
                    result["tables"].append(table.strip())

            # 完整文本
            result["text"] = content

            return result, None

        except Exception as e:
            return None, f"Markdown解析失败: {str(e)}"

    # ========== TXT 加载 ==========
    def _load_txt(self, path: str) -> Tuple[Dict, Optional[str]]:
        """
        加载文本文件，返回结构化信息
        输出格式：
        {
            "type": "report" 或 "log" 或 "unknown",
            "paragraphs": [...],
            "text": "完整文本"
        }
        """
        result = {
            "type": "unknown",
            "paragraphs": [],
            "text": ""
        }

        try:
            with open(path, 'r', encoding='utf-8') as f:
                content = f.read()

            # 按空行分段
            paragraphs = re.split(r'\n\s*\n', content)
            result["paragraphs"] = [p.strip() for p in paragraphs if p.strip()]
            result["text"] = content

            # 类型判断
            sample = content[:500].lower()
            if re.search(r'\d{4}-\d{2}-\d{2}|\d{2}:\d{2}:\d{2}|info|warn|error|debug|trace', sample):
                result["type"] = "log"
            elif re.search(r'第[一二三四五六七八九十]+章|[一二三四]、|摘要|引言|结论|参考文献', sample):
                result["type"] = "report"
            elif re.search(r'公司|企业|合同|协议|条款|甲方|乙方', sample):
                result["type"] = "contract"

            return result, None

        except Exception as e:
            return None, f"文本解析失败: {str(e)}"

    # ========== PDF 加载（新增） ==========
    def _load_pdf(self, path: str) -> Tuple[Dict, Optional[str]]:
        """
        加载PDF文档
        输出格式：
        {
            "type": "pdf",
            "text": "完整文本",
            "pages": [
                {
                    "page_num": 1,
                    "text": "第1页内容",
                    "tables": [markdown_table, ...],
                    "paragraphs": [...]
                }
            ],
            "tables": [markdown_table, ...],
            "paragraphs": [...]
        }
        """
        result = {
            "type": "pdf",
            "pages": [],
            "tables": [],
            "paragraphs": [],
            "text": ""
        }

        # 检查是否安装了pdfplumber
        try:
            import pdfplumber
        except ImportError:
            return None, "请安装 pdfplumber: pip install pdfplumber"

        try:
            with pdfplumber.open(path) as pdf:
                full_text = []

                for page_num, page in enumerate(pdf.pages, 1):
                    # 提取文本
                    page_text = page.extract_text() or ""
                    full_text.append(page_text)

                    # 提取表格
                    page_tables = []
                    tables = page.extract_tables()
                    for table in tables:
                        if table and any(any(cell for cell in row) for row in table):
                            table_md = self._table_to_markdown(table)
                            if table_md:
                                page_tables.append(table_md)
                                result["tables"].append(table_md)

                    # 按行分割作为段落
                    page_paragraphs = []
                    for line in page_text.split('\n'):
                        line = line.strip()
                        if line:
                            page_paragraphs.append(line)
                            result["paragraphs"].append(line)

                    result["pages"].append({
                        "page_num": page_num,
                        "text": page_text,
                        "tables": page_tables,
                        "paragraphs": page_paragraphs
                    })

                result["text"] = "\n".join(full_text)

            return result, None

        except Exception as e:
            return None, f"PDF解析失败: {str(e)}"

    def _table_to_markdown(self, table: List[List]) -> str:
        """
        将表格转换为Markdown格式
        支持从pdfplumber提取的表格
        """
        if not table or not table[0]:
            return ""

        # 清理单元格内容
        cleaned_table = []
        for row in table:
            cleaned_row = []
            for cell in row:
                if cell is None:
                    cleaned_row.append("")
                else:
                    # 移除换行符，替换为空格
                    cell_str = str(cell).replace('\n', ' ').replace('\r', '')
                    # 移除多余空格
                    cell_str = re.sub(r'\s+', ' ', cell_str).strip()
                    cleaned_row.append(cell_str)
            cleaned_table.append(cleaned_row)

        # 过滤掉全空的行
        cleaned_table = [row for row in cleaned_table if any(cell for cell in row)]

        if len(cleaned_table) < 2:  # 至少需要表头+一行数据
            return ""

        # 生成Markdown表格
        md_lines = []

        # 表头
        header = "| " + " | ".join(str(cell) for cell in cleaned_table[0]) + " |"
        md_lines.append(header)

        # 分隔线
        separator = "|" + "|".join(" --- " for _ in cleaned_table[0]) + "|"
        md_lines.append(separator)

        # 数据行
        for row in cleaned_table[1:]:
            # 确保行长度与表头一致
            while len(row) < len(cleaned_table[0]):
                row.append("")
            md_line = "| " + " | ".join(str(cell) for cell in row) + " |"
            md_lines.append(md_line)

        return "\n".join(md_lines)

    # ========== 工具方法 ==========
    def get_info(self, path: str) -> Dict:
        """
        快速获取文件基本信息（不加载内容）
        """
        filename = os.path.basename(path)
        ext = os.path.splitext(filename)[1].lower()
        size = os.path.getsize(path) if os.path.exists(path) else 0

        return {
            "filename": filename,
            "extension": ext,
            "size_bytes": size,
            "size_kb": round(size / 1024, 2),
            "exists": os.path.exists(path)
        }

    def is_supported(self, filename: str) -> bool:
        """
        检查文件格式是否支持
        """
        ext = os.path.splitext(filename)[1].lower()
        return ext in self.supported_formats


# 全局单例
document_loader = DocumentLoader()


# ==================== 便捷函数 ====================

def load_document(path: str) -> Dict:
    """
    快速加载文档
    """
    filename = os.path.basename(path)
    return document_loader.load(path, filename)


def get_loader_info() -> Dict:
    """
    获取加载器信息
    """
    return {
        "supported_formats": document_loader.supported_formats,
        "loader_ready": True
    }